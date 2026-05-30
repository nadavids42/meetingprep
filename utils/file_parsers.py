from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO, StringIO
from pathlib import Path
from typing import BinaryIO

import pandas as pd
from docx import Document
from pypdf import PdfReader


@dataclass
class ParsedSource:
    """Normalized representation of a user-provided source."""

    name: str
    source_type: str
    text: str


SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx", ".pdf", ".csv", ".xlsx"}


def parse_uploaded_file(uploaded_file: BinaryIO) -> ParsedSource:
    """Parse a Streamlit uploaded file into text suitable for LLM extraction."""
    name = getattr(uploaded_file, "name", "uploaded_file")
    extension = Path(name).suffix.lower()
    raw_bytes = uploaded_file.getvalue()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {extension}")

    if extension in {".txt", ".md"}:
        text = _parse_text(raw_bytes)
    elif extension == ".docx":
        text = _parse_docx(raw_bytes)
    elif extension == ".pdf":
        text = _parse_pdf(raw_bytes)
    elif extension == ".csv":
        text = _parse_csv(raw_bytes)
    elif extension == ".xlsx":
        text = _parse_xlsx(raw_bytes)
    else:
        # Defensive fallback; should not be reached because of SUPPORTED_EXTENSIONS.
        raise ValueError(f"Unsupported file type: {extension}")

    return ParsedSource(name=name, source_type=extension.lstrip("."), text=text.strip())


def parse_pasted_notes(notes: str) -> ParsedSource | None:
    """Create a source for pasted context when notes are present."""
    clean_notes = notes.strip()
    if not clean_notes:
        return None
    return ParsedSource(name="Pasted Notes", source_type="text", text=clean_notes)


def _parse_text(raw_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw_bytes.decode("utf-8", errors="replace")


def _parse_docx(raw_bytes: bytes) -> str:
    document = Document(BytesIO(raw_bytes))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    table_lines: list[str] = []
    for table_index, table in enumerate(document.tables, start=1):
        table_lines.append(f"\n[Table {table_index}]")
        for row in table.rows:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            table_lines.append(" | ".join(values))

    return "\n".join(paragraphs + table_lines)


def _parse_pdf(raw_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(raw_bytes))
    page_text: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        page_text.append(f"[Page {page_number}]\n{text.strip()}")
    return "\n\n".join(page_text)


def _parse_csv(raw_bytes: bytes) -> str:
    decoded = _parse_text(raw_bytes)
    dataframe = pd.read_csv(StringIO(decoded))
    return _dataframe_to_context(dataframe, label="CSV")


def _parse_xlsx(raw_bytes: bytes) -> str:
    workbook = pd.read_excel(BytesIO(raw_bytes), sheet_name=None)
    sections: list[str] = []
    for sheet_name, dataframe in workbook.items():
        sections.append(_dataframe_to_context(dataframe, label=f"Sheet: {sheet_name}"))
    return "\n\n".join(sections)


def _dataframe_to_context(dataframe: pd.DataFrame, label: str) -> str:
    """Summarize spreadsheet data without doing complex analytics."""
    row_count = len(dataframe)
    columns = [str(column) for column in dataframe.columns]
    preview = dataframe.fillna("").head(25).to_csv(index=False)

    keyword_pattern = (
        r"(?i)open|blocked|risk|overdue|delayed|issue|action|owner|due|status|uat|"
        r"decision|milestone|dependency|concern"
    )
    relevant_rows = dataframe[
        dataframe.astype(str).apply(
            lambda row: row.str.contains(keyword_pattern, regex=True, na=False).any(), axis=1
        )
    ].head(25)

    relevant_preview = (
        relevant_rows.fillna("").to_csv(index=False)
        if not relevant_rows.empty
        else "No obvious project-tracking rows found by keyword scan."
    )

    return f"""[{label}]
Row count: {row_count}
Column headers: {', '.join(columns)}

Preview rows:
{preview}

Potentially relevant project-tracking rows:
{relevant_preview}
"""
